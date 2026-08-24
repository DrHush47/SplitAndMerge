#!/usr/bin/env python3
"""Programmatic validator for доклад12_structural.docx — checks all 8 structural ПРАВКИ."""

import sys, io, re, zipfile
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FILE = "доклад12_structural.docx"
SRC  = "доклад12.docx"

doc = Document(FILE)
src = Document(SRC)
all_ok = True
oks   = []
warns = []
fails = []

def ok(msg):  oks.append(msg)
def warn(msg): warns.append(msg)
def fail(msg): fails.append(msg); global all_ok; all_ok = False

print(f"=== Programmatic validator — {FILE} ===\n")
print(f"Параграфов: {len(doc.paragraphs)} (исходник: {len(src.paragraphs)})")
print(f"Таблиц:     {len(doc.tables)} (исходник: {len(src.tables)})")

# ═══════════════════════════════════════════════════════════════════════════
# ПРАВКА 1: АКТУАЛЬНОСТЬ before ВВЕДЕНИЕ
# ═══════════════════════════════════════════════════════════════════════════
print("\n=== ПРАВКА 1: АКТУАЛЬНОСТЬ ===")
found_akt = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "АКТУАЛЬНОСТЬ":
        found_akt = True
        bold = all(r.bold for r in p.runs) if p.runs else False
        center = p.alignment == WD_ALIGN_PARAGRAPH.CENTER
        print(f"  P{i}: '{p.text}' (bold={bold}, center={center})")

        # Check АКТУАЛЬНОСТЬ is before ВВЕДЕНИЕ
        vved_idx = None
        for j, q in enumerate(doc.paragraphs):
            if q.text.strip() == "ВВЕДЕНИЕ":
                vved_idx = j; break
        if vved_idx and i < vved_idx:
            print(f"  ✅ АКТУАЛЬНОСТЬ (P{i}) перед ВВЕДЕНИЕ (P{vved_idx})")
            ok("АКТУАЛЬНОСТЬ перед ВВЕДЕНИЕ")
        else:
            print(f"  ❌ АКТУАЛЬНОСТЬ не перед ВВЕДЕНИЕ")
            fail("АКТУАЛЬНОСТЬ не перед ВВЕДЕНИЕ")
        break
if not found_akt:
    print("  ❌ АКТУАЛЬНОСТЬ не найден")
    fail("АКТУАЛЬНОСТЬ не найден")

# ═══════════════════════════════════════════════════════════════════════════
# ПРАВКА 2: НОВИЗНА between АКТУАЛЬНОСТЬ and ВВЕДЕНИЕ
# ═══════════════════════════════════════════════════════════════════════════
print("\n=== ПРАВКА 2: НОВИЗНА ===")
found_nov = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "НОВИЗНА":
        found_nov = True
        print(f"  P{i}: '{p.text}'")
        akt_idx = vved_idx = None
        for j, q in enumerate(doc.paragraphs):
            if q.text.strip() == "АКТУАЛЬНОСТЬ": akt_idx = j
            if q.text.strip() == "ВВЕДЕНИЕ": vved_idx = j
        if akt_idx is not None and vved_idx is not None and akt_idx < i < vved_idx:
            print(f"  ✅ Порядок: АКТУАЛЬНОСТЬ (P{akt_idx}) → НОВИЗНА (P{i}) → ВВЕДЕНИЕ (P{vved_idx})")
            ok("НОВИЗНА в правильном порядке")
        else:
            print(f"  ❌ Порядок нарушен (akt={akt_idx}, nov={i}, vved={vved_idx})")
            fail("НОВИЗНА не в правильном порядке")
        break
if not found_nov:
    print("  ❌ НОВИЗНА не найден")
    fail("НОВИЗНА не найден")

# ═══════════════════════════════════════════════════════════════════════════
# ПРАВКА 3: Структурированная аннотация
# ═══════════════════════════════════════════════════════════════════════════
print("\n=== ПРАВКА 3: Структурированная аннотация ===")
found_struct = False
for i, p in enumerate(doc.paragraphs):
    if "Структурированная аннотация" in p.text:
        found_struct = True
        print(f"  P{i}: '{p.text[:80]}'")
        has_cel = any(q.text.startswith("Цель исследования") for q in doc.paragraphs)
        has_mat = any(q.text.startswith("Материалы и методы") for q in doc.paragraphs)
        has_rez = any(q.text.startswith("Результаты") for q in doc.paragraphs)
        has_vyv = any(q.text.startswith("Выводы") for q in doc.paragraphs)
        print(f"  Цель={has_cel}, Материалы={has_mat}, Результаты={has_rez}, Выводы={has_vyv}")
        if has_cel and has_mat and has_rez and has_vyv:
            print(f"  ✅ Все 4 секции аннотации есть")
            ok("Структурированная аннотация: 4 секции")
        else:
            missing = []
            if not has_cel: missing.append("Цель")
            if not has_mat: missing.append("Материалы")
            if not has_rez: missing.append("Результаты")
            if not has_vyv: missing.append("Выводы")
            print(f"  ❌ Не хватает: {missing}")
            fail(f"Структурированная аннотация: нет секций {missing}")

        # Check original "Аннотация." is gone
        has_old_annot = any(q.text.strip().startswith("Аннотация.") for q in doc.paragraphs)
        if not has_old_annot:
            print(f"  ✅ Старый параграф 'Аннотация.' удалён")
            ok("Старый 'Аннотация.' удалён")
        else:
            print(f"  ❌ Старый 'Аннотация.' всё ещё в документе")
            fail("Старый 'Аннотация.' не удалён")
        break
if not found_struct:
    print("  ❌ Структурированная аннотация не найдена")
    fail("Структурированная аннотация отсутствует")

# ═══════════════════════════════════════════════════════════════════════════
# ПРАВКА 4: ORCID
# ═══════════════════════════════════════════════════════════════════════════
print("\n=== ПРАВКА 4: ORCID ===")
orcid_ps = [(i, p) for i, p in enumerate(doc.paragraphs) if "ORCID" in p.text]
orcid_count = len(orcid_ps)
print(f"  ORCID параграфов: {orcid_count}")
for i, p in orcid_ps:
    print(f"    P{i}: '{p.text[:80]}'")
if orcid_count == 3:
    print(f"  ✅ 3 ORCID")
    ok("3 ORCID строки")
elif orcid_count == 0:
    print(f"  ⚠️ ORCID не найдены")
    warn("ORCID не найдены")
else:
    print(f"  ⚠️ Ожидалось 3, фактически {orcid_count}")
    warn(f"ORCID: ожидалось 3, получено {orcid_count}")

# ═══════════════════════════════════════════════════════════════════════════
# ПРАВКА 5: Ключевые слова нумерованные
# ═══════════════════════════════════════════════════════════════════════════
print("\n=== ПРАВКА 5: Ключевые слова нумерованные ===")
found_kw = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Ключевые слова:":
        found_kw = True
        numbered = 0
        nums = []
        for j in range(i+1, min(i+15, len(doc.paragraphs))):
            q = doc.paragraphs[j]
            m = re.match(r'^(\d+)\.\s', q.text.strip())
            if m:
                numbered += 1
                nums.append(int(m.group(1)))
            elif numbered > 0:
                break
        print(f"  P{i}: 'Ключевые слова:' + {numbered} нумерованных пунктов → {nums}")
        if numbered >= 5:
            print(f"  ✅ {numbered} нумерованных пунктов")
            ok(f"Ключевые слова: {numbered} пунктов")
        else:
            print(f"  ❌ Только {numbered} пунктов (ожидалось ≥5)")
            fail(f"Ключевые слова: только {numbered} пунктов")
        break
if not found_kw:
    print("  ❌ 'Ключевые слова:' не найдено")
    fail("Ключевые слова не найдены")

# ═══════════════════════════════════════════════════════════════════════════
# ПРАВКА 6: Колонтитул с номером страницы
# ═══════════════════════════════════════════════════════════════════════════
print("\n=== ПРАВКА 6: Колонтитул ===")
try:
    section = doc.sections[0]
    header = section.header
    header_text = " ".join(p.text for p in header.paragraphs)
    print(f"  Header text: '{header_text[:80]}'")

    # Check XML for PAGE field
    has_page_field = False
    if header.paragraphs:
        hxml = header.paragraphs[0]._element.xml
        has_page_field = "PAGE" in hxml or "fldChar" in hxml

    if has_page_field or "Страница" in header_text:
        print(f"  ✅ Колонтитул с номером страницы")
        ok("Колонтитул с PAGE")
    else:
        print(f"  ❌ Колонтитул не настроен (xml has fldChar: {'fldChar' in hxml if header.paragraphs else 'N/A'})")
        fail("Колонтитул не настроен")
except Exception as e:
    print(f"  ❌ Ошибка: {e}")
    fail(f"Ошибка колонтитула: {e}")

# ═══════════════════════════════════════════════════════════════════════════
# ПРАВКА 7: Гиперссылка DOI в списке литературы
# ═══════════════════════════════════════════════════════════════════════════
print("\n=== ПРАВКА 7: Гиперссылка DOI ===")
ref_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "ЛИТЕРАТУРА":
        ref_start = i + 1; break
if ref_start:
    # Find first non-empty ref
    first_ref = None
    for i in range(ref_start, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if p.text.strip():
            first_ref = p; break

    if first_ref:
        xml_str = first_ref._element.xml
        has_link = "hyperlink" in xml_str.lower() or "doi.org" in first_ref.text.lower()
        print(f"  Первый источник: '{first_ref.text[:80]}'")
        print(f"  Hyperlink in XML: {has_link}")
        if has_link:
            print(f"  ✅ Гиперссылка добавлена")
            ok("DOI гиперссылка")
        else:
            print(f"  ⚠️ Гиперссылка не найдена в XML (проверить в Word)")
            warn("Гиперссылка не подтверждена")
    else:
        print("  ❌ Нет ссылок после ЛИТЕРАТУРА")
        fail("Список литературы пуст")
else:
    print("  ❌ ЛИТЕРАТУРА не найдена")
    fail("ЛИТЕРАТУРА не найдена")

# ═══════════════════════════════════════════════════════════════════════════
# ПРАВКА 8: КОНФЛИКТ ИНТЕРЕСОВ before ЛИТЕРАТУРА
# ═══════════════════════════════════════════════════════════════════════════
print("\n=== ПРАВКА 8: КОНФЛИКТ ИНТЕРЕСОВ ===")
found_konfl = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "КОНФЛИКТ ИНТЕРЕСОВ":
        found_konfl = True
        lit_idx = None
        for j, q in enumerate(doc.paragraphs):
            if q.text.strip() == "ЛИТЕРАТУРА":
                lit_idx = j; break
        if lit_idx and i < lit_idx:
            print(f"  ✅ КОНФЛИКТ ИНТЕРЕСОВ (P{i}) перед ЛИТЕРАТУРА (P{lit_idx})")
            ok("КОНФЛИКТ ИНТЕРЕСОВ перед ЛИТЕРАТУРА")
        else:
            print(f"  ❌ Не перед ЛИТЕРАТУРА (konfl={i}, lit={lit_idx})")
            fail("КОНФЛИКТ ИНТЕРЕСОВ не перед ЛИТЕРАТУРА")
        break
if not found_konfl:
    print("  ❌ КОНФЛИКТ ИНТЕРЕСОВ не найден")
    fail("КОНФЛИКТ ИНТЕРЕСОВ не найден")

# ═══════════════════════════════════════════════════════════════════════════
# ИТОГ
# ═══════════════════════════════════════════════════════════════════════════
print("\n" + "=" * 64)
print(f"\n✅ OK:   {len(oks)}")
for m in oks: print(f"   ✅ {m}")
print(f"\n⚠️  WARN: {len(warns)}")
for m in warns: print(f"   ⚠️  {m}")
print(f"\n❌ FAIL: {len(fails)}")
for m in fails: print(f"   ❌ {m}")
print("\n" + "=" * 64)
status = "✅ ВСЕ КРИТИЧЕСКИЕ ПРОВЕРКИ ПРОЙДЕНЫ" if not fails else "❌ ЕСТЬ ПРОБЛЕМЫ"
print(f"  ИТОГ: {status}")
print(f"  Всего OK: {len(oks)}/{len(oks) + len(warns) + len(fails)}")
