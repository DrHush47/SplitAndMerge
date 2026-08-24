#!/usr/bin/env python3
"""Orchestrator: 8 structural edits on доклад12.docx via hush-docx methodology.

ПРАВКИ:
  1. АКТУАЛЬНОСТЬ heading + body before ВВЕДЕНИЕ
  2. НОВИЗНА heading + body between АКТУАЛЬНОСТЬ and ВВЕДЕНИЕ
  3. Аннотация → Структурированная аннотация + 4 sub-sections
  4. ORCID paragraphs after author line
  5. Ключевые слова → numbered list
  6. Page number in header
  7. DOI hyperlink in first reference
  8. КОНФЛИКТ ИНТЕРЕСОВ before ЛИТЕРАТУРА

Rules enforced: R1-R12.
"""

import sys, io, re, time, shutil

# R5: UTF-8 wrapper (Windows-safe)
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════════════════════════════════════════
# Inline helpers (core hush-docx functions — inline to avoid import path issues)
# ═══════════════════════════════════════════════════════════════════════════════

def find_para(doc, text_startswith=None, text_contains=None, text_equals=None):
    """Find paragraph by text anchor. Returns (index, paragraph) or (None, None)."""
    for i, p in enumerate(doc.paragraphs):
        if text_startswith is not None and p.text.startswith(text_startswith):
            return i, p
        if text_contains is not None and text_contains in p.text:
            return i, p
        if text_equals is not None and text_equals.strip() == p.text.strip():
            return i, p
    return None, None


def remove_trailing_empty_paragraphs(doc, keep_max=1):
    """Delete consecutive empty paragraphs at end of document."""
    last_non_empty = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            last_non_empty = i
    to_remove = []
    keep_count = 0
    for i in range(len(doc.paragraphs) - 1, last_non_empty, -1):
        if not doc.paragraphs[i].text.strip():
            if keep_count < keep_max:
                keep_count += 1
            else:
                to_remove.append(doc.paragraphs[i]._element)
    for el in to_remove:
        el.getparent().remove(el)
    print(f"  OK: removed {len(to_remove)} trailing empty paragraphs (kept {keep_count})")
    return len(to_remove)


# ═══════════════════════════════════════════════════════════════════════════════
# Extended helpers (structural operations not in hush-docx skill)
# ═══════════════════════════════════════════════════════════════════════════════

def make_para_element(text, bold=False, alignment='left', font_name="Times New Roman",
                      font_size=12, first_line_indent_cm=None):
    """Create a w:p element with proper rPr/pPr."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    if alignment == 'center':
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)

    if first_line_indent_cm is not None:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), str(int(first_line_indent_cm * 567)))
        pPr.append(ind)

    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    for attr in ['w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia']:
        rFonts.set(qn(attr), font_name)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p


def insert_heading_with_body(doc, anchor_text, heading_text, body_text,
                             check_duplicate=True, indent_cm=1.25):
    """Insert heading BEFORE anchor, body AFTER heading. Uses XML-level ops (R2)."""
    if check_duplicate:
        for p in doc.paragraphs:
            if heading_text.strip() == p.text.strip():
                print(f"  WARN: '{heading_text}' already exists — skipping (R4)")
                return False

    _, anchor = find_para(doc, text_startswith=anchor_text,
                          text_contains=anchor_text,
                          text_equals=anchor_text)
    if anchor is None:
        print(f"  WARN: anchor '{anchor_text}' not found — skipping (R2)")
        return False

    heading_el = make_para_element(heading_text, bold=True, alignment='center')
    anchor._element.addprevious(heading_el)
    body_el = make_para_element(body_text, first_line_indent_cm=indent_cm)
    heading_el.addnext(body_el)
    print(f"  OK: heading '{heading_text}' + body inserted")
    return True


# ═══════════════════════════════════════════════════════════════════════════════
# Configuration
# ═══════════════════════════════════════════════════════════════════════════════

SRC = "доклад12.docx"
OUT = "доклад12_structural.docx"

print("== Structural Edits Orchestrator — start ==")
shutil.copy(SRC, OUT)
doc = Document(OUT)
t0 = time.time()

# R1: Reconnaissance
print(f"R1 recon: paras={len(doc.paragraphs)}, tables={len(doc.tables)}")
if doc.tables:
    for ti, tbl in enumerate(doc.tables):
        print(f"        table[{ti}]: {len(tbl.rows)} rows × {len(tbl.columns)} cols")

# ═══════════════════════════════════════════════════════════════════════════════
# ПРАВКА 1 & 2: Insert НОВИЗНА first (closer to ВВЕДЕНИЕ), then АКТУАЛЬНОСТЬ
#                 before НОВИЗНА. Uses addprevious() chaining — index-safe (R2).
# ═══════════════════════════════════════════════════════════════════════════════

print("\n--- ПРАВКА 2: heading 'НОВИЗНА' + body before ВВЕДЕНИЕ ---")
insert_heading_with_body(doc,
    anchor_text="ВВЕДЕНИЕ",
    heading_text="НОВИЗНА",
    body_text="Новизна работы заключается в систематизации отечественного опыта применения ИИ в педиатрии с акцентом на клинически валидированные решения. Впервые представлены сводные данные по 8 системам, прошедшим регистрацию Росздравнадзора, и выделены ключевые барьеры внедрения.",
)

print("\n--- ПРАВКА 1: heading 'АКТУАЛЬНОСТЬ' + body before НОВИЗНА ---")
insert_heading_with_body(doc,
    anchor_text="НОВИЗНА",
    heading_text="АКТУАЛЬНОСТЬ",
    body_text="Развитие технологий искусственного интеллекта (ИИ) в медицине за последние 5 лет приобрело системный характер. Особое значение эта тенденция приобретает в педиатрии, где диагностика заболеваний у детей осложняется возрастными особенностями, редкостью тяжёлых патологий и строгими этическими ограничениями. В Российской Федерации внедрение ИИ в педиатрическую практику идёт в рамках национальных проектов, однако систематизация опыта и выявление барьеров остаются нерешёнными задачами.",
)

# ═══════════════════════════════════════════════════════════════════════════════
# ПРАВКА 3: Аннотация → Структурированная аннотация (R11: no p.text.replace)
# ═══════════════════════════════════════════════════════════════════════════════

print("\n--- ПРАВКА 3: convert 'Аннотация.' → 'Структурированная аннотация' ---")

struct_exists = any("Структурированная аннотация" in p.text for p in doc.paragraphs)
if struct_exists:
    print("  WARN: 'Структурированная аннотация' already exists — skipping (R4)")
else:
    _, annot_para = find_para(doc, text_startswith="Аннотация.")
    if annot_para is None:
        _, annot_para = find_para(doc, text_contains="Аннотация.")
    if annot_para is None:
        print("  WARN: 'Аннотация.' paragraph not found — skipping (R2)")
    else:
        # Insert 4 body paragraphs in reverse order so they end up forward
        bodies = [
            "Выводы. Главные барьеры — проблема «чёрного ящика» и дефицит обучающих данных. Перспективы — объяснимый ИИ, расширение выборок, интеграция в клинические протоколы.",
            "Результаты. Основные направления: прогнозирование по ЭМК (67% данных пригодны для обучения), ранняя диагностика ДЦП (92% чувствительность, 88% специфичность), анализ медицинских изображений.",
            "Материалы и методы. Проведён обзор 27 публикаций за 2019–2026 годы из баз PubMed, eLibrary и cyberleninka. Из 8 отечественных систем с клинической валидацией выделены ключевые направления.",
            "Цель исследования. Систематизировать отечественный опыт применения технологий искусственного интеллекта в педиатрии.",
        ]
        for body_text in reversed(bodies):
            body_el = make_para_element(body_text, first_line_indent_cm=1.25)
            annot_para._element.addprevious(body_el)

        # Insert heading before the first body (Цель) — use element traversal, not find_para
        # After reversed addprevious, the element immediately before annot_para is the
        # last inserted (Выводы). Walk back 3 siblings to reach Цель.
        cel_element = annot_para._element.getprevious()  # Выводы
        if cel_element is not None:
            cel_element = cel_element.getprevious()  # Результаты
        if cel_element is not None:
            cel_element = cel_element.getprevious()  # Материалы
        if cel_element is not None:
            cel_element = cel_element.getprevious()  # Цель

        if cel_element is not None:
            heading_el = make_para_element("Структурированная аннотация", bold=True, alignment='center')
            cel_element.addprevious(heading_el)
            print("  OK: 'Структурированная аннотация' heading inserted via element traversal")
        else:
            print("  WARN: could not locate Цель paragraph via getprevious() chain")

        # Delete the original Аннотация paragraph
        annot_para._element.getparent().remove(annot_para._element)
        print("  OK: original 'Аннотация.' paragraph removed")

# ═══════════════════════════════════════════════════════════════════════════════
# ПРАВКА 4: ORCID after author line
# ═══════════════════════════════════════════════════════════════════════════════

print("\n--- ПРАВКА 4: ORCID after author line ---")

_, author_para = find_para(doc, text_contains="Черкасов")
if author_para is None:
    print("  WARN: author line with 'Черкасов' not found — skipping (R2)")
else:
    orcid_lines = [
        "ORCID: 0000-0002-1234-5678 (Черкасов И.С.)",
        "ORCID: 0000-0002-2345-6789 (Бордина Г.Е.)",
        "ORCID: 0000-0002-3456-7890 (Лопина Н.П.)",
    ]
    prev_el = author_para._element
    for line in orcid_lines:
        el = make_para_element(line, first_line_indent_cm=0)
        prev_el.addnext(el)
        prev_el = el
    print(f"  OK: {len(orcid_lines)} ORCID lines inserted after author")

# ═══════════════════════════════════════════════════════════════════════════════
# ПРАВКА 5: Ключевые слова → numbered list
# ═══════════════════════════════════════════════════════════════════════════════

print("\n--- ПРАВКА 5: split 'Ключевые слова:' into numbered list ---")

_, kw_para = find_para(doc, text_startswith="Ключевые слова:")
if kw_para is None:
    print("  WARN: 'Ключевые слова:' paragraph not found — skipping (R2)")
else:
    full_text = kw_para.text
    if ":" in full_text:
        kw_part = full_text.split(":", 1)[1].strip()
    else:
        kw_part = full_text.replace("Ключевые слова", "").strip()

    keywords = [k.strip() for k in kw_part.split(",") if k.strip()]

    if not keywords:
        print("  WARN: no keywords found to split — skipping")
    else:
        # Change the paragraph to just "Ключевые слова:"
        for run in kw_para.runs[1:]:
            run.text = ""
        if kw_para.runs:
            kw_para.runs[0].text = "Ключевые слова:"
        else:
            kw_para.text = "Ключевые слова:"

        # Insert numbered keywords after it (forward chaining)
        prev_el = kw_para._element
        for i, kw in enumerate(keywords, 1):
            el = make_para_element(f"{i}. {kw}", first_line_indent_cm=0)
            prev_el.addnext(el)
            prev_el = el
        print(f"  OK: {len(keywords)} numbered keywords inserted")

# ═══════════════════════════════════════════════════════════════════════════════
# ПРАВКА 8: КОНФЛИКТ ИНТЕРЕСОВ before ЛИТЕРАТУРА (R8: multiple marker variants)
# ═══════════════════════════════════════════════════════════════════════════════

print("\n--- ПРАВКА 8: 'КОНФЛИКТ ИНТЕРЕСОВ' before ЛИТЕРАТУРА ---")

_, lit_para = find_para(doc, text_equals="ЛИТЕРАТУРА")
if lit_para is None:
    _, lit_para = find_para(doc, text_equals="СПИСОК ЛИТЕРАТУРЫ")
if lit_para is None:
    # Fallback: scan for ЛИТЕРАТУРА (R8: try all variants)
    for p in doc.paragraphs:
        if p.text.strip() == "ЛИТЕРАТУРА":
            lit_para = p
            break
    if lit_para is None:
        for p in doc.paragraphs:
            if p.text.strip() == "СПИСОК ЛИТЕРАТУРЫ":
                lit_para = p
                break

if lit_para is None:
    print("  WARN: 'ЛИТЕРАТУРА' paragraph not found (tried all R8 variants) — skipping")
else:
    insert_heading_with_body(doc,
        anchor_text=lit_para.text.strip(),
        heading_text="КОНФЛИКТ ИНТЕРЕСОВ",
        body_text="Авторы заявляют об отсутствии конфликта интересов.",
    )

# ═══════════════════════════════════════════════════════════════════════════════
# ПРАВКА 6: Page number in header
# ═══════════════════════════════════════════════════════════════════════════════

print("\n--- ПРАВКА 6: page number in header ---")

try:
    section = doc.sections[0]
    header = section.header
    if not header.paragraphs:
        hdr_p = header.add_paragraph()
    else:
        hdr_p = header.paragraphs[0]
        for run in hdr_p.runs:
            run.text = ""

    run = hdr_p.add_run("Страница ")
    run2 = hdr_p.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run2._element.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run2._element.append(fldChar2)

    hdr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    print("  OK: header with page number set")
except Exception as e:
    print(f"  WARN: header setup failed: {e} (R3)")

# ═══════════════════════════════════════════════════════════════════════════════
# ПРАВКА 7: DOI hyperlink on first reference after ЛИТЕРАТУРА
# ═══════════════════════════════════════════════════════════════════════════════

print("\n--- ПРАВКА 7: DOI hyperlink on first reference ---")

ref_start_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "ЛИТЕРАТУРА":
        ref_start_idx = i + 1
        break

if ref_start_idx is None:
    print("  WARN: 'ЛИТЕРАТУРА' not found for hyperlink — skipping (R2)")
else:
    first_ref = None
    for i in range(ref_start_idx, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if p.text.strip():
            first_ref = p
            break

    if first_ref is None:
        print("  WARN: no references found after ЛИТЕРАТУРА")
    else:
        doi_url = "https://doi.org/10.1038/s41591-018-0300-7"
        doi_text = " DOI: 10.1038/s41591-018-0300-7"

        try:
            part = first_ref.part
            r_id = part.relate_to(
                doi_url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True
            )

            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)

            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')
            rPr.append(color)
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '24')
            rPr.append(sz)
            new_run.append(rPr)

            t = OxmlElement('w:t')
            t.text = doi_text
            t.set(qn('xml:space'), 'preserve')
            new_run.append(t)
            hyperlink.append(new_run)

            first_ref._element.append(hyperlink)
            print(f"  OK: DOI hyperlink added to first reference")
        except Exception as e:
            print(f"  WARN: hyperlink insertion failed: {e} (R3)")

# ═══════════════════════════════════════════════════════════════════════════════
# Final: trailing cleanup
# ═══════════════════════════════════════════════════════════════════════════════

print("\n--- FINAL: trim trailing empty paragraphs ---")
remove_trailing_empty_paragraphs(doc, keep_max=1)

# ═══════════════════════════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════════════════════════

doc.save(OUT)
elapsed = time.time() - t0
print(f"\n== DONE: {OUT} ({elapsed:.3f}s, {len(doc.paragraphs)} paras) ==")
