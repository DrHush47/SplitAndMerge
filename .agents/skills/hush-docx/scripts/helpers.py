"""
Battle-tested python-docx helper functions for routine mechanical .docx edits.

Each function encodes an empirical rule (R1-R12) discovered in real tests.
See SKILL.md for the full rule set and rationale.

Usage:
    from scripts.helpers import set_run_font, set_para_format, insert_heading_before, ...
"""

import re
from docx.shared import Pt, Mm, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── R5-compatible: font that survives Cyrillic ──────────────────────────────

def set_run_font(run, name="Times New Roman", size=12, bold=None, italic=None):
    """Set font on a run, including cs & eastAsia attributes (R5: critical for Cyrillic)."""
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ['w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia']:
        rFonts.set(qn(attr), name)


# ── Paragraph & cell formatting ─────────────────────────────────────────────

def set_para_format(para, alignment=None, first_line_indent_cm=None,
                    line_spacing=None, bold=None,
                    font_name="Times New Roman", font_size=12):
    """Apply paragraph-level formatting and propagate font to all child runs."""
    if alignment is not None:
        para.alignment = alignment
    if first_line_indent_cm is not None:
        para.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    if line_spacing is not None:
        para.paragraph_format.line_spacing = line_spacing
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    for run in para.runs:
        set_run_font(run, font_name, font_size, bold=bold)


def set_cell_format(cell, font_name="Times New Roman", font_size=12):
    """Apply table-cell formatting: zero indent, single spacing, font."""
    for para in cell.paragraphs:
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        for run in para.runs:
            set_run_font(run, font_name, font_size)


# ── Paragraph search (R2: text anchors, not hardcoded indexes) ──────────────

def find_para(doc, text_startswith=None, text_contains=None, text_equals=None):
    """Find a paragraph by text anchor. Returns (index, paragraph) or (None, None)."""
    for i, p in enumerate(doc.paragraphs):
        if text_startswith is not None and p.text.startswith(text_startswith):
            return i, p
        if text_contains is not None and text_contains in p.text:
            return i, p
        if text_equals is not None and text_equals.strip() == p.text.strip():
            return i, p
    return None, None


# ── Heading insertion (R4: check_duplicate=True always) ─────────────────────

def insert_heading_before(doc, anchor_text, heading_text, check_duplicate=True):
    """Insert a bold, centered, TNR-12pt heading before the anchor paragraph.

    Returns True if inserted, False if skipped (duplicate / missing anchor).
    """
    if check_duplicate:
        for i, p in enumerate(doc.paragraphs):
            if heading_text.strip() in p.text:
                print(f"  WARN: '{heading_text}' already exists at P{i} — skipping (R4)")
                return False
    _, anchor = find_para(doc, text_startswith=anchor_text)
    if anchor is None:
        print(f"  WARN: anchor '{anchor_text}' not found — skipping (R2)")
        return False

    new_p = OxmlElement('w:p')
    new_pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    new_pPr.append(jc)
    new_p.append(new_pPr)

    new_r = OxmlElement('w:r')
    new_rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    for attr in ['w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia']:
        rFonts.set(qn(attr), 'Times New Roman')
    new_rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 12 pt
    new_rPr.append(sz)
    b = OxmlElement('w:b')
    new_rPr.append(b)
    new_r.append(new_rPr)
    t = OxmlElement('w:t')
    t.text = heading_text
    new_r.append(t)
    new_p.append(new_r)

    anchor._element.addprevious(new_p)
    print(f"  OK: heading '{heading_text}' inserted")
    return True


# ── End-of-document insertion (R10: do LAST among insertion edits) ──────────

def add_paragraph_safe(doc, text, check_duplicate_text=True):
    """Append a paragraph with duplicate guard.

    Returns True if added, False if skipped (duplicate detected).
    """
    if check_duplicate_text:
        first_line = text.split('\n')[0][:80]
        for p in doc.paragraphs:
            if first_line in p.text:
                print(f"  WARN: text block '{first_line}...' already exists — skipping (R10)")
                return False
    doc.add_paragraph(text)
    print("  OK: paragraph added to end")
    return True


# ── Cross-run text replacement (table cells, etc.) ──────────────────────────

def replace_cross_run_text(paragraph, old, new):
    """Replace a substring across run boundaries.

    WARNING: discards formatting of runs[1:].  Not suitable for paragraphs
    with heterogeneous formatting across runs.
    """
    if not paragraph.runs:
        return False
    full_text = paragraph.text
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    for run in paragraph.runs[1:]:
        run.text = ""
    paragraph.runs[0].text = new_text
    return True


def replace_cross_run_in_cell(cell, old, new):
    """Replace a substring in every paragraph of a table cell (cross-run aware)."""
    found = False
    for para in cell.paragraphs:
        if replace_cross_run_text(para, old, new):
            found = True
    return found


# ── Reference renumbering (R7 / R8 / R12) ───────────────────────────────────

def renumber_references(doc, start_patterns, stop_patterns):
    """Renumber references between start_patterns and stop_patterns as 1..N.

    start_patterns: list of header texts (e.g. ['ЛИТЕРАТУРА', 'REFERENCES']).
    stop_patterns:  list of EXPLICIT section markers (R12: never regex on body text).

    Strips existing numbering with regex '^\\s*\\d+\\.\\s*' before renumbering.
    Returns the number of references processed.
    """
    ref_start = None
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        for sp in start_patterns:
            if txt == sp:
                ref_start = i + 1
                break
        if ref_start is not None:
            break

    if ref_start is None:
        print(f"  WARN: no start header among {start_patterns}")
        return 0

    n = 1
    for i in range(ref_start, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        txt = p.text.strip()
        if not txt:
            continue
        if txt in stop_patterns:
            break
        cleaned = re.sub(r'^\s*\d+\.\s*', '', txt)
        new_text = f"{n}. {cleaned}"
        for run in p.runs[1:]:
            run.text = ""
        if p.runs:
            p.runs[0].text = new_text
        else:
            p.text = new_text
        n += 1

    return n - 1


# ── Trailing cleanup ────────────────────────────────────────────────────────

def remove_trailing_empty_paragraphs(doc, keep_max=1):
    """Delete consecutive empty paragraphs at the end of the document.

    keep_max: number of trailing empty paragraphs to preserve (default 1).
    Returns the number of paragraphs removed.
    """
    last_non_empty = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            last_non_empty = i

    to_remove = []
    keep_count = 0
    for i in range(len(doc.paragraphs) - 1, last_non_empty, -1):
        if not doc.paragraphs[i].text.strip():
            if keep_count < keep_max:
                keep_count += 1
            else:
                to_remove.append(doc.paragraphs[i]._element)

    for el in to_remove:
        el.getparent().remove(el)

    print(f"  OK: removed {len(to_remove)} trailing empty paragraphs (kept {keep_count})")
    return len(to_remove)
